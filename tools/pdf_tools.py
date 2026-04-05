# -*- coding: utf-8 -*-
import os
from fpdf import FPDF   # 用 fpdf2 避免 reportlab md5 错误
import markdown
import subprocess
from pdf2docx import Converter
import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 初始化变量
weasyprint = None
WEASYPRINT_AVAILABLE = False

# 尝试导入 weasyprint
if False:  # 暂时禁用 weasyprint，避免 cairo 依赖问题
    try:
        import weasyprint
        WEASYPRINT_AVAILABLE = True
    except ImportError:
        print("Warning: WeasyPrint not installed, will use FPDF as fallback")
else:
    print("Warning: WeasyPrint disabled, will use FPDF as fallback")

def txt_to_pdf(txt_path, pdf_path, font_size=12):
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    
    # 首先尝试使用 weasyprint（如果可用）
    if WEASYPRINT_AVAILABLE:
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 转换为 HTML，添加适当的样式
            html_content = f'''
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        font-size: {font_size}px;
                        margin: 20px;
                        line-height: 1.5;
                    }}
                    pre {{
                        white-space: pre-wrap;
                        word-wrap: break-word;
                        font-family: Arial, sans-serif;
                        margin: 0;
                    }}
                </style>
            </head>
            <body>
                <pre>{content}</pre>
            </body>
            </html>
            '''
            
            # 使用 weasyprint 生成 PDF
            weasyprint.HTML(string=html_content).write_pdf(pdf_path)
            return True
        except Exception as e:
            print(f"WeasyPrint 转换失败: {e}")
    
    # 如果 weasyprint 不可用或失败，使用 FPDF
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=font_size)
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # 读取文本并处理长行
        with open(txt_path, 'r', encoding='utf-8') as f:
            for line in f:
                text = line.rstrip()
                if not text:
                    pdf.ln(5)
                    continue
                
                # 处理长行，手动换行
                max_width = pdf.w - 20  # 页面宽度减去边距
                words = text.split()
                current_line = []
                current_width = 0
                
                for word in words:
                    word_width = pdf.get_string_width(word) + pdf.get_string_width(' ')
                    if current_width + word_width <= max_width:
                        current_line.append(word)
                        current_width += word_width
                    else:
                        # 写入当前行
                        if current_line:
                            pdf.cell(0, 10, ' '.join(current_line))
                            pdf.ln()
                        # 开始新行
                        current_line = [word]
                        current_width = word_width
                
                # 写入最后一行
                if current_line:
                    pdf.cell(0, 10, ' '.join(current_line))
                    pdf.ln()
        
        pdf.output(pdf_path)
        return True
    except Exception as e:
        print(f"FPDF 转换失败: {e}")
        return False

def markdown_to_pdf(md_path, pdf_path):
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    if WEASYPRINT_AVAILABLE:
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                html = markdown.markdown(f.read())
            weasyprint.HTML(string=html).write_pdf(pdf_path)
            return True
        except Exception as e:
            print(f"Markdown 转 PDF 失败: {e}")
    
    # 如果 weasyprint 不可用或失败，使用 FPDF 作为备选
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # 读取 Markdown 文件内容
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 简单处理 Markdown，将其转换为纯文本
        lines = content.split('\n')
        for line in lines:
            if not line:
                pdf.ln(5)
                continue
            
            # 处理长行
            max_width = pdf.w - 20  # 页面宽度减去边距
            words = line.split()
            current_line = []
            current_width = 0
            
            for word in words:
                word_width = pdf.get_string_width(word) + pdf.get_string_width(' ')
                if current_width + word_width <= max_width:
                    current_line.append(word)
                    current_width += word_width
                else:
                    # 写入当前行
                    if current_line:
                        pdf.cell(0, 10, ' '.join(current_line))
                        pdf.ln()
                    # 开始新行
                    current_line = [word]
                    current_width = word_width
            
            # 写入最后一行
            if current_line:
                pdf.cell(0, 10, ' '.join(current_line))
                pdf.ln()
        
        pdf.output(pdf_path)
        print("使用 FPDF 成功转换 Markdown 为 PDF")
        return True
    except Exception as e:
        print(f"FPDF 转换 Markdown 失败: {e}")
        return False

def latex_to_pdf(tex_path, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    cmd = ["pdflatex", "-interaction=nonstopmode", "-output-directory", output_folder, tex_path]
    try:
        subprocess.run(cmd, check=True)
        return True
    except Exception as e:
        print(f"LaTeX 转 PDF 失败: {e}")
        return False

def fast_pdf_to_docx(pdf_path, docx_path):
    """快速转换扫描版PDF到DOCX（仅提取图像）"""
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    try:
        # 打开PDF并提取图像
        pdf = fitz.open(pdf_path)
        num_pages = len(pdf)
        print(f"开始快速转换扫描版PDF，共{num_pages}页")
        
        # 创建DOCX文档
        doc = Document()
        
        # 设置页面大小为A4，无页边距
        section = doc.sections[0]
        # A4页面大小：210mm x 297mm
        section.page_width = Pt(595)  # A4宽度（点）
        section.page_height = Pt(842)  # A4高度（点）
        section.left_margin = Pt(0)  # 无边距
        section.right_margin = Pt(0)  # 无边距
        section.top_margin = Pt(0)  # 无边距
        section.bottom_margin = Pt(0)  # 无边距
        
        # 使用整个页面宽度
        available_width = section.page_width
        
        # 处理第一页（添加提示信息和第一页图像）
        if num_pages > 0:
            # 提取第一页图像
            page = pdf.load_page(0)
            pix = page.get_pixmap(dpi=300)  # 设置300DPI，提高清晰度
            img_path = f"temp_page_0.png"
            pix.save(img_path)
            
            # 添加红色大字提示（作为图像的标题）
            para = doc.add_paragraph()
            run = para.add_run("扫描版PDF没有文字层")
            run.font.size = Pt(28)  # 减小字体大小，避免占用太多空间
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
            
            # 添加简短说明
            para2 = doc.add_paragraph()
            para2.add_run("此文档由扫描版PDF转换而来，仅包含原始图像。")
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加第一页图像
            doc.add_picture(img_path, width=available_width)  # 使用整个页面宽度
            
            # 删除临时图像
            if os.path.exists(img_path):
                os.remove(img_path)
            
            # 移除分页符，让内容连续排列
            # if num_pages > 1:
            #     doc.add_page_break()
        
        # 处理剩余页面
        for page_num in range(1, num_pages):
            if page_num % 5 == 0:
                print(f"处理第{page_num+1}/{num_pages}页")
            
            page = pdf.load_page(page_num)
            # 提取页面为高清晰度图像
            pix = page.get_pixmap(dpi=300)  # 设置300DPI，提高清晰度
            # 保存临时图像
            img_path = f"temp_page_{page_num}.png"
            pix.save(img_path)
            
            # 将图像添加到DOCX
            doc.add_picture(img_path, width=available_width)  # 使用整个页面宽度
            
            # 移除分页符，让内容连续排列
            # if page_num < num_pages - 1:
            #     doc.add_page_break()  # 添加分页符，确保每页一个图像
            
            # 删除临时图像
            if os.path.exists(img_path):
                os.remove(img_path)
        
        pdf.close()
        doc.save(docx_path)
        doc = None  # 释放内存
        print(f"快速转换完成，保存为: {docx_path}")
        return True
    except Exception as e:
        print(f"快速转换失败: {e}")
        return False

def get_libreoffice_cmd():
    import platform
    system = platform.system()
    if system == "Windows":
        return r"C:\Program Files\LibreOffice\program\soffice.exe"
    else:
        return "libreoffice"

LIBREOFFICE_CMD = get_libreoffice_cmd()

def pdf_to_docx(pdf_path, docx_path):
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    
    # 优先尝试使用LibreOffice
    try:
        # 确保路径是绝对路径，避免LibreOffice路径问题
        pdf_path = os.path.abspath(pdf_path)
        output_folder = os.path.abspath(os.path.dirname(docx_path))
        
        # 构建转换命令
        cmd = [LIBREOFFICE_CMD, "--headless", "--convert-to", "docx", "--outdir", output_folder, pdf_path]
        
        # 执行命令并捕获输出
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            # LibreOffice会生成与PDF同名的DOCX文件
            libreoffice_output = os.path.join(output_folder, f"{os.path.splitext(os.path.basename(pdf_path))[0]}.docx")
            if os.path.exists(libreoffice_output):
                # 如果生成的文件与目标路径不同，重命名
                if libreoffice_output != docx_path:
                    if os.path.exists(docx_path):
                        os.remove(docx_path)
                    os.rename(libreoffice_output, docx_path)
                print(f"使用LibreOffice成功转换PDF为DOCX: {docx_path}")
                return True
            else:
                print("LibreOffice转换失败：未生成输出文件")
        else:
            print(f"LibreOffice转换失败: {result.stderr}")
    except Exception as e:
        print(f"LibreOffice转换失败: {e}")
    
    # 如果LibreOffice失败，尝试使用pdf2docx库
    try:
        # 先检查PDF是否为扫描版（Words count: 0）
        doc = fitz.open(pdf_path)
        total_words = 0
        # 只检查前3页，进一步减少检查时间
        for page_num in range(min(3, len(doc))):
            page = doc.load_page(page_num)
            text = page.get_text()
            total_words += len(text.split())
        
        # 检查PDF页数和大小
        num_pages = len(doc)
        file_size = os.path.getsize(pdf_path) / (1024 * 1024)  # 转换为MB
        doc.close()  # 及时关闭文件，释放内存
        
        is_scanned = total_words == 0
        
        # 处理文件名：如果是扫描版，在文件名前添加标记
        if is_scanned:
            base_dir = os.path.dirname(docx_path)
            base_name = os.path.basename(docx_path)
            new_base_name = f"（扫描版PDF没有文字层）{base_name}"
            docx_path = os.path.join(base_dir, new_base_name)
            # 对于扫描版PDF，使用快速转换方法
            success = fast_pdf_to_docx(pdf_path, docx_path)
            if not success:
                # 扫描版PDF转换失败，给出明确的警告
                print(f"⚠️ 警告：扫描版PDF转换失败，请检查文件是否损坏")
            return success
        else:
            # 对于可转化为文字的PDF，使用完整的转换参数
            print(f"使用完整转换模式 ({num_pages}页, {file_size:.2f}MB)，包含布局保持、表格分析、结构分析和OCR")
            cv = Converter(pdf_path)
            cv.convert(docx_path, 
                      start=0, 
                      end=None, 
                      pages=None, 
                      keep_layout=True,  # 保持布局
                      zoom=1.0, 
                      table_floating=True,  # 启用浮动表格分析
                      use_ocr=True,  # 启用OCR
                      parse_structure=True)  # 启用结构分析
            cv.close()  # 及时关闭转换器，释放内存
            return True
    except Exception as e:
        # 检查是否为扫描版PDF
        try:
            doc = fitz.open(pdf_path)
            total_words = 0
            for page_num in range(min(3, len(doc))):
                page = doc.load_page(page_num)
                text = page.get_text()
                total_words += len(text.split())
            doc.close()
            is_scanned = total_words == 0
            if is_scanned:
                print(f"⚠️ 警告：扫描版PDF转换失败，错误：{e}")
            else:
                print(f"PDF → Word失败: {e}")
        except:
            print(f"PDF → Word失败: {e}")
        return False