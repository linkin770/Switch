# -*- coding: utf-8 -*-
import os
import platform
import subprocess
from config import CONFIG

# 尝试导入python-office相关库
poexcel_available = False
poppt_available = False
poword_available = False
pandas_available = False
python_docx_available = False
weasyprint_available = False
comtypes_available = False
win32com_available = False

try:
    import poexcel
    poexcel_available = True
except ImportError:
    pass
try:
    import poppt
    poppt_available = True
except ImportError:
    pass
try:
    import poword
    poword_available = True
except ImportError:
    pass
try:
    import pandas as pd
    pandas_available = True
except ImportError:
    pass
try:
    from docx import Document
    python_docx_available = True
except ImportError:
    pass
try:
    import weasyprint
    weasyprint_available = True
except ImportError:
    pass
try:
    import comtypes.client
    comtypes_available = True
except ImportError:
    pass
try:
    import win32com.client
    import pythoncom
    win32com_available = True
except ImportError:
    pass

# 使用config.py中的LibreOffice路径
LIBREOFFICE_CMD = CONFIG.LIBREOFFICE_CMD


def _try_libreoffice(input_path: str, output_folder: str, pdf_path: str) -> str | None:
    """尝试使用 LibreOffice 转换"""
    try:
        abs_input = os.path.abspath(input_path)
        abs_output = os.path.abspath(output_folder)
        cmd = [LIBREOFFICE_CMD, "--headless", "--convert-to", "pdf", "--outdir", abs_output, abs_input]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0 and os.path.exists(pdf_path):
            print(f"LibreOffice转换成功: {input_path}")
            return pdf_path
        else:
            print(f"LibreOffice转换失败: {result.stderr.strip() or '未知错误'}")
    except Exception as e:
        print(f"LibreOffice转换失败: {e}")
    return None


def _try_comtypes_excel(input_path: str, pdf_path: str) -> str | None:
    """尝试使用 comtypes (Excel COM) 转换"""
    if not comtypes_available:
        return None
    try:
        import comtypes.client
        print("尝试使用 comtypes (Excel COM) 转换")
        xl_app = comtypes.client.CreateObject("Excel.Application")
        xl_app.Visible = False
        xl_app.DisplayAlerts = 0
        books = xl_app.Workbooks.Open(os.path.abspath(input_path), False)
        books.ExportAsFixedFormat(0, os.path.abspath(pdf_path))
        books.Close(False)
        xl_app.Quit()
        if os.path.exists(pdf_path):
            print(f"comtypes Excel转换成功: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"comtypes Excel转换失败: {e}")
    return None


def _try_comtypes_powerpoint(input_path: str, pdf_path: str) -> str | None:
    """尝试使用 comtypes (PowerPoint COM) 转换"""
    if not comtypes_available:
        return None
    try:
        import comtypes.client
        print("尝试使用 comtypes (PowerPoint COM) 转换")
        ppt_app = comtypes.client.CreateObject("PowerPoint.Application")
        ppt_app.Visible = True
        presentation = ppt_app.Presentations.Open(os.path.abspath(input_path), WithWindow=False)
        presentation.SaveAs(os.path.abspath(pdf_path), 32)  # 32 = ppSaveAsPDF
        presentation.Close()
        ppt_app.Quit()
        if os.path.exists(pdf_path):
            print(f"comtypes PowerPoint转换成功: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"comtypes PowerPoint转换失败: {e}")
    return None


def _try_comtypes_word(input_path: str, pdf_path: str) -> str | None:
    """尝试使用 comtypes (Word COM) 转换"""
    if not comtypes_available:
        return None
    try:
        import comtypes.client
        print("尝试使用 comtypes (Word COM) 转换")
        word_app = comtypes.client.CreateObject("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0
        doc = word_app.Documents.Open(os.path.abspath(input_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word_app.Quit()
        if os.path.exists(pdf_path):
            print(f"comtypes Word转换成功: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"comtypes Word转换失败: {e}")
    return None


def _try_win32com_word(input_path: str, pdf_path: str) -> str | None:
    """尝试使用 win32com (Word) 转换"""
    if not win32com_available:
        return None
    try:
        import win32com.client
        import pythoncom
        print("尝试使用 win32com (Word) 转换")
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(input_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        if os.path.exists(pdf_path):
            print(f"win32com Word转换成功: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"win32com Word转换失败: {e}")
    return None


def _try_win32com_excel(input_path: str, pdf_path: str) -> str | None:
    """尝试使用 win32com (Excel) 转换"""
    if not win32com_available:
        return None
    try:
        import win32com.client
        import pythoncom
        print("尝试使用 win32com (Excel) 转换")
        pythoncom.CoInitialize()
        xl = win32com.client.Dispatch('Excel.Application')
        xl.Visible = False
        xl.DisplayAlerts = 0
        wb = xl.Workbooks.Open(os.path.abspath(input_path))
        wb.ExportAsFixedFormat(0, os.path.abspath(pdf_path))
        wb.Close(False)
        xl.Quit()
        pythoncom.CoUninitialize()
        if os.path.exists(pdf_path):
            print(f"win32com Excel转换成功: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"win32com Excel转换失败: {e}")
    return None


def _try_win32com_powerpoint(input_path: str, pdf_path: str) -> str | None:
    """尝试使用 win32com (PowerPoint) 转换"""
    if not win32com_available:
        return None
    try:
        import win32com.client
        import pythoncom
        print("尝试使用 win32com (PowerPoint) 转换")
        pythoncom.CoInitialize()
        ppt = win32com.client.Dispatch('PowerPoint.Application')
        presentation = ppt.Presentations.Open(os.path.abspath(input_path), WithWindow=False)
        presentation.SaveAs(os.path.abspath(pdf_path), 32)  # 32 = ppSaveAsPDF
        presentation.Close()
        ppt.Quit()
        pythoncom.CoUninitialize()
        if os.path.exists(pdf_path):
            print(f"win32com PowerPoint转换成功: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"win32com PowerPoint转换失败: {e}")
    return None


def _find_generated_pdf(output_folder: str, base_name: str, pdf_path: str) -> str | None:
    """检查期望路径或同目录下的同名PDF是否存在"""
    if os.path.exists(pdf_path):
        return pdf_path
    alt = os.path.join(output_folder, f"{base_name}.pdf")
    if os.path.exists(alt) and alt != pdf_path:
        return alt
    return None


def convert_office_to_pdf(input_path: str, output_folder: str) -> str | None:
    os.makedirs(output_folder, exist_ok=True)

    ext = os.path.splitext(input_path)[1].lower()
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    pdf_path = os.path.join(output_folder, f"{base_name}.pdf")

    # ========== DOCX 回退链 ==========
    if ext == ".docx":
        # 方法1: python-docx + weasyprint
        if python_docx_available and weasyprint_available:
            try:
                print("[DOCX → PDF] 方法1: python-docx + weasyprint")
                doc = Document(input_path)
                html_content = '''<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
body { font-family: Arial, sans-serif; font-size: 12px; margin: 20px; line-height: 1.5; }
h1,h2,h3,h4,h5,h6 { margin-top: 20px; margin-bottom: 10px; }
p { margin-bottom: 10px; }
table { border-collapse: collapse; width: 100%; margin-bottom: 10px; }
th,td { border: 1px solid #ddd; padding: 8px; text-align: left; }
th { background-color: #f2f2f2; }
ul,ol { margin-bottom: 10px; }
</style></head><body>'''
                for paragraph in doc.paragraphs:
                    text = paragraph.text
                    if text:
                        style_name = paragraph.style.name
                        if style_name.startswith('Heading'):
                            level = style_name.split(' ')[1] if len(style_name.split(' ')) > 1 else '1'
                            html_content += f'<h{level}>{text}</h{level}>'
                        else:
                            html_content += f'<p>{text}</p>'
                    else:
                        html_content += '<p>&nbsp;</p>'
                for table in doc.tables:
                    html_content += '<table>'
                    for row in table.rows:
                        html_content += '<tr>'
                        for cell in row.cells:
                            html_content += f'<td>{cell.text}</td>'
                        html_content += '</tr>'
                    html_content += '</table>'
                html_content += '</body></html>'
                weasyprint.HTML(string=html_content).write_pdf(pdf_path)
                result = _find_generated_pdf(output_folder, base_name, pdf_path)
                if result:
                    print(f"[DOCX → PDF] 方法1成功: {result}")
                    return result
            except Exception as e:
                print(f"[DOCX → PDF] 方法1失败: {e}")

        # 方法2: poword
        if poword_available:
            try:
                print("[DOCX → PDF] 方法2: poword")
                poword.docx2pdf(path=input_path, output_path=os.path.dirname(pdf_path))
                result = _find_generated_pdf(output_folder, base_name, pdf_path)
                if result:
                    print(f"[DOCX → PDF] 方法2成功: {result}")
                    return result
            except Exception as e:
                print(f"[DOCX → PDF] 方法2失败: {e}")

        # 方法3: comtypes Word COM
        result = _try_comtypes_word(input_path, pdf_path)
        if result:
            return result

        # 方法4: LibreOffice
        result = _try_libreoffice(input_path, output_folder, pdf_path)
        if result:
            return result

        # 方法5: win32com Word
        result = _try_win32com_word(input_path, pdf_path)
        if result:
            return result

    # ========== XLS / XLSX 回退链 ==========
    elif ext in [".xls", ".xlsx"]:
        # 方法1: poexcel
        if poexcel_available:
            try:
                print("[Excel → PDF] 方法1: poexcel")
                try:
                    poexcel.excel2pdf(input_path, os.path.dirname(pdf_path))
                except TypeError:
                    try:
                        poexcel.excel2pdf(input_path=input_path, output_path=os.path.dirname(pdf_path))
                    except TypeError:
                        poexcel.excel2pdf(file=input_path, output=os.path.dirname(pdf_path))
                result = _find_generated_pdf(output_folder, base_name, pdf_path)
                if result:
                    print(f"[Excel → PDF] 方法1成功: {result}")
                    return result
            except AttributeError:
                print("[Excel → PDF] 方法1: poexcel缺少excel2pdf函数")
            except Exception as e:
                print(f"[Excel → PDF] 方法1失败: {e}")

        # 方法2: comtypes Excel COM
        result = _try_comtypes_excel(input_path, pdf_path)
        if result:
            return result

        # 方法3: LibreOffice
        result = _try_libreoffice(input_path, output_folder, pdf_path)
        if result:
            return result

        # 方法4: win32com Excel
        result = _try_win32com_excel(input_path, pdf_path)
        if result:
            return result

    # ========== PPT / PPTX 回退链 ==========
    elif ext in [".ppt", ".pptx"]:
        # 方法1: poppt
        if poppt_available:
            try:
                print("[PPT → PDF] 方法1: poppt")
                poppt.ppt2pdf(path=input_path, output_path=os.path.dirname(pdf_path))
                result = _find_generated_pdf(output_folder, base_name, pdf_path)
                if result:
                    print(f"[PPT → PDF] 方法1成功: {result}")
                    return result
            except Exception as e:
                print(f"[PPT → PDF] 方法1失败: {e}")

        # 方法2: comtypes PowerPoint COM
        result = _try_comtypes_powerpoint(input_path, pdf_path)
        if result:
            return result

        # 方法3: LibreOffice
        result = _try_libreoffice(input_path, output_folder, pdf_path)
        if result:
            return result

        # 方法4: win32com PowerPoint
        result = _try_win32com_powerpoint(input_path, pdf_path)
        if result:
            return result

    # ========== 其他格式（DOC, RTF, HTML, XML）回退链 ==========
    else:
        # 方法1: LibreOffice（通用转换器）
        result = _try_libreoffice(input_path, output_folder, pdf_path)
        if result:
            return result

        # 方法2: comtypes Word COM（仅 DOC/RTF）
        if ext in [".doc", ".rtf"]:
            result = _try_comtypes_word(input_path, pdf_path)
            if result:
                return result
            result = _try_win32com_word(input_path, pdf_path)
            if result:
                return result

    print(f"[转换失败] 所有方法均无法将 {input_path} 转换为 PDF")
    return None


def csv_to_xlsx(csv_path: str, xlsx_path: str) -> bool:
    os.makedirs(os.path.dirname(xlsx_path), exist_ok=True)
    try:
        if pandas_available:
            pd.read_csv(csv_path).to_excel(xlsx_path, index=False)
            return True
        else:
            print("pandas不可用，无法转换CSV到Excel")
            return False
    except Exception as e:
        print(f"CSV → Excel失败: {e}")
        return False

def xlsx_to_csv(xlsx_path: str, csv_path: str) -> bool:
    os.makedirs(os.path.dirname(csv_path), exist_ok=True)
    try:
        if pandas_available:
            # 读取Excel文件
            df = pd.read_excel(xlsx_path)
            
            # 检查是否有Unnamed列
            has_unnamed = any('Unnamed' in col for col in df.columns)
            
            # 如果有Unnamed列，为其生成默认列名
            if has_unnamed:
                # 生成默认列名，如A, B, C, ..., Z, AA, AB, ...
                def generate_column_name(index):
                    name = ''
                    while index >= 0:
                        name = chr(65 + index % 26) + name
                        index = index // 26 - 1
                    return name
                
                # 为所有列生成新的列名
                new_columns = [generate_column_name(i) for i in range(len(df.columns))]
                df.columns = new_columns
            
            # 保存为CSV文件
            df.to_csv(csv_path, index=False)
            return True
        else:
            print("pandas不可用，无法转换Excel到CSV")
            return False
    except Exception as e:
        print(f"Excel → CSV失败: {e}")
        return False
